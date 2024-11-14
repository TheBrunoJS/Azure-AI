import requests, os
from docx import Document
from utils.config import Config


def translate_text(text, target_language):

    path = '/translate'
    constructed_url = Config.ENDPOINT+path
    
    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }

    headers = {
        'Ocp-Apim-Subscription-Key': Config.KEY,
        # location required if you're using a multi-service or regional (not global) resource.
        'Ocp-Apim-Subscription-Region': Config.LOCATION,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{
    'text': text
    }]

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    
    return response[0]["translations"][0]["text"]


def translate_document(path, target_language):
    document = Document(path)
    full_text = []
    translated_document = Document()
    
    for paragraph in document.paragraphs:
        translated_text = translate_text(paragraph.text, target_language)
        full_text.append(translated_text)
    for line in full_text:
        translated_document.add_paragraph(line)
    
    path_translated = path.replace(".docx", f"_{target_language}.docx")
    translated_document.save(path_translated)

    return path_translated